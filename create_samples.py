import os
import fitz
import docx
import pandas as pd
from pathlib import Path

def create_samples():
    sample_dir = Path('sample_files')
    sample_dir.mkdir(exist_ok=True)

    # 1. TXT - Meeting Minutes
    with open(sample_dir / 'meeting_minutes.txt', 'w', encoding='utf-8') as f:
        f.write("Weekly Team Meeting - Jan 15, 2026\n\n")
        f.write("Attendees: Alice, Bob, Charlie\n")
        f.write("Agenda:\n- Q1 Roadmap Review\n- Budget Allocation\n- Office Renovation Update\n\n")
        f.write("Action Items:\n1. Alice to finalize the roadmap by Friday.\n2. Bob to send budget report.\n")

    # 2. DOCX - Project Proposal
    doc = docx.Document()
    doc.add_heading('Project Phoenix Proposal', 0)
    doc.add_paragraph('Executive Summary', style='Heading 1')
    doc.add_paragraph('Project Phoenix aims to revolutionize our legacy systems by migrating to a cloud-native architecture. This will improve scalability and reduce maintenance costs by 40%.')
    doc.add_paragraph('Key Objectives', style='Heading 1')
    doc.add_paragraph('1. Zero downtime deployment.\n2. Enhanced security compliance.\n3. Real-time analytics dashboard.')
    doc.save(sample_dir / 'project_proposal.docx')

    # 3. XLSX - Q1 Budget
    data = {
        'Category': ['Software Licenses', 'Hardware Upgrades', 'Team Training', 'Marketing', 'Total'],
        'Budget': [5000, 12000, 3000, 8000, 28000],
        'Spent': [4500, 2000, 500, 1500, 8500],
        'Remaining': [500, 10000, 2500, 6500, 19500]
    }
    df = pd.DataFrame(data)
    df.to_excel(sample_dir / 'q1_budget.xlsx', index=False)

    # 4. PDF - Invoice
    doc = fitz.open()
    page = doc.new_page()
    
    # Simple layout simulation
    page.insert_text((50, 50), "INVOICE #2026-001", fontsize=20)
    page.insert_text((50, 80), "Date: January 13, 2026")
    page.insert_text((50, 100), "Bill To: Acme Corp")
    
    y = 150
    page.insert_text((50, y), "Description", fontsize=12)
    page.insert_text((300, y), "Amount", fontsize=12)
    
    items = [
        ("Consulting Services (10 hrs)", "$1,500.00"),
        ("System Audit", "$800.00"),
        ("Report Generation", "$200.00")
    ]
    
    for desc, amt in items:
        y += 20
        page.insert_text((50, y), desc)
        page.insert_text((300, y), amt)
        
    page.insert_text((300, y + 40), "Total: $2,500.00", fontsize=14, color=(0, 0, 1))
    
    doc.save(sample_dir / 'invoice_2026_001.pdf')
    
    print(f"Sample files created in {sample_dir.resolve()}")

if __name__ == '__main__':
    create_samples()
